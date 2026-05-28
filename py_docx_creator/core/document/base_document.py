from pathlib import Path
from typing import Callable, Any

from py_docx_creator.abstract_classes.abc_document.abc_base_document import ABCBaseDocument
from docx import Document as DocxDocument  # alias

from py_docx_creator.abstract_classes.abc_document.abc_document import ABCDocument


class BaseDocument(ABCBaseDocument):
    """
    Класс билдера Документа

    Attributes:
        path (Path | str | None): Путь до документа
        name (str | None): Наименование документа
        document (DocxDocument): Класс документа python-docx
    """

    path: Path | str | None = None  # путь до документа
    name: str | None = None  # наименование документа
    _creation_instruction: Callable | None = None  # инструкция для формирования документа
    _instruction_kwargs: dict[str, Any] | None  # аргументы инструкция для формирования документа
    document: type[DocxDocument] | None = None  # alias

    def __init__(self, file_name: str | None = None, directory: str | Path | None = None):
        if file_name is None:
            return

        directory = Path(directory) if directory else Path.cwd()
        full_path = directory / file_name

        self.path = full_path
        self.name = file_name

        if self._is_valid_existing_file(full_path):
            self.load_document()
        else:
            self.create_document()

    @staticmethod
    def _is_valid_existing_file(path: Path) -> bool:
        return not path.name.startswith("~") and path.suffix == ".docx" and path.exists()

    def create_document(self):
        self.document = DocxDocument()

    def load_document(self):
        self.document = DocxDocument(self.path)

    def save_document(self):
        if not self.document:
            raise ValueError("Документ не инициализирован")
        self.document.save(self.path)

    def run_instruction(self, document: "ABCDocument"):
        if self.creation_instruction:
            self.creation_instruction(document, **self._instruction_kwargs)
        else:
            raise Exception("Инструкция не задана!")

    @property
    def creation_instruction(self) -> Callable:
        return self._creation_instruction

    @creation_instruction.setter
    def creation_instruction(self, value: Callable) -> None:
        self._creation_instruction = value

    @property
    def instruction_kwargs(self) -> dict[str | Any] | None:
        return self._instruction_kwargs

    @instruction_kwargs.setter
    def instruction_kwargs(self, value: dict[str | Any] | None) -> None:
        self._instruction_kwargs = value
