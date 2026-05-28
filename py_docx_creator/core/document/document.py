from pathlib import Path
from typing import Any, Self, Callable

from docx import Document as DocxDocument  # alias
from docx.styles.styles import Styles
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from py_docx_creator.abstract_classes.abc_document.abc_base_document import ABCBaseDocument
from py_docx_creator.abstract_classes.abc_document.abc_document import ABCDocument
from py_docx_creator.abstract_classes.abc_document.abc_document_style import ABCDocumentStyle
from py_docx_creator.abstract_classes.abc_document.abc_document_writer import ABCDocumentWriter
from py_docx_creator.abstract_classes.abc_style_dataclasses.abc_paragraph_style import ABCParagraphStyle
from py_docx_creator.abstract_classes.abc_style_dataclasses.abc_text_style import ABCTextStyle
from py_docx_creator.core.document.base_document import BaseDocument
from py_docx_creator.core.document.builder import Builder
from py_docx_creator.core.document.document_style import DocumentStyle
from py_docx_creator.core.document.document_writer import DocumentWriter
from py_docx_creator.enums.enum_align_paragraph import AlignParagraph


class Document(ABCDocument):
    """Абстрактный класс фасада (User Document API)"""

    _writer: ABCDocumentWriter
    _style: ABCDocumentStyle
    _doc_builder: ABCBaseDocument

    def __init__(self, file_name: str | Path | None = None, path: str | Path | None = None):
        self._doc_builder = BaseDocument(file_name=file_name, directory=path)
        self._writer = DocumentWriter(self)
        self._style = DocumentStyle()

    @property
    def document(self) -> DocxDocument:
        return self._doc_builder.document

    @document.setter
    def document(self, value: DocxDocument) -> None:
        self._doc_builder.document._document = value

    @property
    def name(self):
        return self._doc_builder.name

    @name.setter
    def name(self, value: str) -> None:
        self._doc_builder.name = value

    def add_paragraph_to_document(self) -> Paragraph:
        return self._writer.add_paragraph_to_document(self)

    def add_run_to_paragraph(self, paragraph: Paragraph, text: str) -> Run | None:
        return self._writer.add_run_to_paragraph(paragraph, text)

    def add_page_break(self) -> None:
        self._writer.add_page_break(self)

    def write(
        self,
        text: str,
        paragraph_style: ABCParagraphStyle | type[ABCParagraphStyle],
        text_style: ABCTextStyle | type[ABCTextStyle],
        target: Paragraph | None = None,
        size: float | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
        underline: bool | None = None,
        space_after: float | None = None,
        alignment: AlignParagraph | None = None,
        first_line_indent: float | None = None,
        with_leader: bool = False,
        leader_width: float = 6.8,
    ) -> Paragraph:
        """
        Метод быстрой записи в документ

        Args:
            target (Document | Paragraph): Цель записи
            text (str): Записываемый текст
            paragraph_style (ABCParagraphStyle): Стиль создаваемого параграфа
            text_style (ABCTextStyle): Стиль записываемого текста
            size (float | None): Размер шрифта записываемого текста
            bold (bool | None): Включить ли жирное начертание для записываемого текста
            italic (bool | None): Включить ли курсивное начертание для записываемого текста
            underline (bool | None): Включить ли подчеркнутое начертание для записываемого текста
            space_after (float | None): Размер отступа после созданного параграфа
            alignment (AlignParagraph | None): Выравнивание для созданного параграфа
            first_line_indent (float | None): Размер отступа для первой строки (красная строка)
            with_leader (bool | None): Включить ли заполнение остатка строки табуляцией со знаком '_'
            leader_width (float): Размер заполняемой табуляцией строки. По умолчанию: 6.8
        Returns:
            Paragraph: Созданный параграф
        """
        kwargs = locals().copy()
        kwargs.pop("self")
        return self._writer.write(**kwargs)

    def paragraph(self, text: str) -> Builder:
        return self._writer.paragraph(text)

    @property
    def document_style(self) -> str | None:
        """Стиль документа"""
        return self._style.document_style

    @document_style.setter
    def document_style(self, value: str) -> None:
        """Стиль документа"""
        self._style.document_style = value

    def get_document_style(self) -> Styles:
        return self._style.get_document_style(self)

    def apply_style(self, target: Self | Run | Paragraph, style: Any) -> None:
        """
        Применение стиля к передаваемому объекту

        Args:
            target (ABCDocument | Run | Paragraph): Цель применения стиля
            style (Any): Применяемый стиль
        """
        self._style.apply_style(target, style)

    def create_document(self, file_name: str, path: str | Path | None) -> None:
        """
        Создание документа

        Args:
            file_name (str): Наименование документа
            path (str | Path | None): Путь формирования документа
        """
        self._doc_builder.create_document(file_name, path)

    def load_document(self) -> None:
        """Загрузка уже имеющегося документа"""
        self._doc_builder.load_document()

    def save_document(self) -> None:
        """Сохранение документа"""
        self._doc_builder.save_document()

    def run_instruction(self) -> None:
        """
        Запуск инструкции формирования документа
        """
        self._doc_builder.run_instruction()

    @property
    def creation_instruction(self) -> Callable:
        """Функция для формирования документа"""
        return self._doc_builder.creation_instruction

    @creation_instruction.setter
    def creation_instruction(self, value: Callable) -> None:
        """Функция для формирования документа"""
        self._doc_builder.creation_instruction = value

    @property
    def instruction_kwargs(self) -> dict[str | Any] | None:
        return self._doc_builder.instruction_kwargs

    @instruction_kwargs.setter
    def instruction_kwargs(self, value: dict[str | Any] | None) -> None:
        self._doc_builder.instruction_kwargs = value
