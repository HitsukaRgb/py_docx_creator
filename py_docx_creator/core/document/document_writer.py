from copy import copy
from typing import Type

from docx.enum.text import WD_TAB_LEADER, WD_TAB_ALIGNMENT
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from py_docx_creator.abstract_classes.abc_document.abc_document import ABCDocument
from py_docx_creator.abstract_classes.abc_document.abc_document_writer import ABCDocumentWriter
from py_docx_creator.abstract_classes.abc_style_dataclasses.abc_paragraph_style import ABCParagraphStyle
from py_docx_creator.abstract_classes.abc_style_dataclasses.abc_text_style import ABCTextStyle
from py_docx_creator.core.document.document_style import DocumentStyle
from py_docx_creator.enums.enum_align_paragraph import AlignParagraph
from docx import Document as DocxDocument  # alias

from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from py_docx_creator.core.document.document import Document


class DocumentWriter(ABCDocumentWriter):
    """
    Класс для наполнения документа

    Attributes:
        document (DocxDocument): Класс документа python-docx
    """
    document: DocxDocument  # alias

    def add_paragraph_to_document(self, document: "Document | None" = None) -> Paragraph | None:
        return document.document.add_paragraph() if document else self.document.add_paragraph()

    @staticmethod
    def add_run_to_paragraph(paragraph: Paragraph, text: str) -> Run | None:
        return paragraph.add_run(text)

    def add_page_break(self, document: "Document | None" = None) -> None:
        document.document.add_page_break() if document else self.document.add_page_break()


class FastWriter(DocumentWriter):
    """Класс быстрой записи в документ"""

    def write(self,
              target: "Document | Paragraph",
              text: str,
              paragraph_style: ABCParagraphStyle | Type[ABCParagraphStyle],
              text_style: ABCTextStyle | Type[ABCTextStyle],
              size: float | None = None,
              bold: bool | None = None,
              italic: bool | None = None,
              underline: bool | None = None,
              space_after: float | None = None,
              alignment: AlignParagraph | None = None,
              first_line_indent: float | None = None,
              with_leader: bool = False,
              leader_width: float = 6.8
              ) -> Paragraph:
        """
        Метод быстрой записи в документ

        Args:
            target (Document | Paragraph): Цель записи
            text (str): Записываемый текст
            paragraph_style (ABCParagraphStyle): Стиль создаваемого параграфа
            text_style (ABCTextStyle): Стиль записываемого текста
            size (float | None): Размер шрифта записываемого текста
            bold (bool | None): Включить ли жирное начертание для записываемого текста
            italic (bool | None): Включить ли курсивное начертание для записываемого текста
            underline (bool | None): Включить ли подчеркнутое начертание для записываемого текста
            space_after (float | None): Размер отступа после созданного параграфа
            alignment (AlignParagraph | None): Выравнивание для созданного параграфа
            first_line_indent (float | None): Размер отступа для первой строки (красная строка)
            with_leader (bool | None): Включить ли заполнение остатка строки табуляцией со знаком '_'
            leader_width (float): Размер заполняемой табуляцией строки. По умолчанию: 6.8
        Returns:
            Paragraph: Созданный параграф
        """
        if any(val is not None for val in [bold, italic, underline, size, alignment, first_line_indent, space_after]):
            paragraph_style = copy(paragraph_style)
            text_style = copy(text_style)

            if bold is not None: text_style.bold = bold
            if italic is not None: text_style.italic = italic
            if underline is not None: text_style.underline = underline
            if size is not None: text_style.size = size
            if alignment is not None: paragraph_style.alignment = alignment.value
            if first_line_indent is not None: paragraph_style.first_line_indent = first_line_indent
            if space_after is not None: paragraph_style.space_after = space_after

        if isinstance(target, ABCDocument):
            # При передаче документа в качестве цели для записи
            paragraph = target.add_paragraph_to_document(target)
        else:
            # В случае если целью для записи является параграф
            paragraph = target

        if with_leader:
            text = f"{text}\t"
            tab_stops = paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(
                Inches(leader_width),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
                leader=WD_TAB_LEADER.LINES
            )

        DocumentStyle.apply_style(paragraph, paragraph_style)
        run = self.add_run_to_paragraph(paragraph, text)
        DocumentStyle.apply_style(run, text_style)
        return paragraph


class Writer(FastWriter):
    """Класс писателя"""
    pass
