from abc import ABC, abstractmethod
from pathlib import Path
from typing import Callable, Any, Self
from docx import Document as DocxDocument  # alias
from docx.styles.styles import Styles
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from py_docx_creator.abstract_classes.abc_document.abc_base_document import ABCBaseDocument
from py_docx_creator.abstract_classes.abc_document.abc_document_style import ABCDocumentStyle
from py_docx_creator.abstract_classes.abc_document.abc_document_writer import ABCDocumentWriter
from py_docx_creator.abstract_classes.abc_style_dataclasses.abc_paragraph_style import ABCParagraphStyle
from py_docx_creator.abstract_classes.abc_style_dataclasses.abc_text_style import ABCTextStyle
from py_docx_creator.core.document.builder import Builder
from py_docx_creator.enums.enum_align_paragraph import AlignParagraph


class ABCDocument(ABC):
    """Абстрактный класс фасада (User Document API)"""

    _writer: ABCDocumentWriter
    _style: ABCDocumentStyle
    _doc_builder: ABCBaseDocument

    @abstractmethod
    def __init__(self, file_name: str | Path | None = None, path: str | Path | None = None):
        """
        Конструктор
        Arguments:
            file_name (str | Path | None): Наименование документа
            path (str | Path | None): Директория формирования документа
        """
        pass

    @property
    @abstractmethod
    def document(self) -> DocxDocument:
        """Класс документа python-docx"""
        pass

    @document.setter
    def document(self, value: DocxDocument) -> None:
        pass

    @abstractmethod
    def add_paragraph_to_document(self) -> Paragraph:
        """Добавление параграфа в документ"""
        pass

    @staticmethod
    @abstractmethod
    def add_run_to_paragraph(paragraph: Paragraph, text: str) -> Run | None:
        """
        Добавить текст в параграф

        Args:
            paragraph (Paragraph): Параграф для записи Run
            text (str): Текст записываемый в Run
        Return:
            Run | None: Возвращает записанный Run или None
        """
        pass

    @abstractmethod
    def add_page_break(self) -> None:
        """Добавление разрыва страницы в документ"""
        pass

    @abstractmethod
    def write(
        self,
        text: str,
        paragraph_style: ABCParagraphStyle | type[ABCParagraphStyle],
        text_style: ABCTextStyle | type[ABCTextStyle],
        target: "Paragraph | None" = None,
        size: float | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
        underline: bool | None = None,
        space_after: float | None = None,
        alignment: AlignParagraph | None = None,
        first_line_indent: float | None = None,
        with_leader: bool = False,
        leader_width: float = 6.8,
    ) -> Paragraph:
        """
        Метод быстрой записи в документ

        Args:
            target (Document | Paragraph): Цель записи
            text (str): Записываемый текст
            paragraph_style (ABCParagraphStyle): Стиль создаваемого параграфа
            text_style (ABCTextStyle): Стиль записываемого текста
            size (float | None): Размер шрифта записываемого текста
            bold (bool | None): Включить ли жирное начертание для записываемого текста
            italic (bool | None): Включить ли курсивное начертание для записываемого текста
            underline (bool | None): Включить ли подчеркнутое начертание для записываемого текста
            space_after (float | None): Размер отступа после созданного параграфа
            alignment (AlignParagraph | None): Выравнивание для созданного параграфа
            first_line_indent (float | None): Размер отступа для первой строки (красная строка)
            with_leader (bool | None): Включить ли заполнение остатка строки табуляцией со знаком '_'
            leader_width (float): Размер заполняемой табуляцией строки. По умолчанию: 6.8
        Returns:
            Paragraph: Созданный параграф
        """
        pass

    @abstractmethod
    def paragraph(self, text: str) -> Builder:
        """
        Инициализация style билдера

        Args:
            text (str | None): Записываемый в документ текст
        Returns:
            Builder: Экземпляр style билдера
        """
        pass

    @property
    @abstractmethod
    def document_style(self) -> str | None:
        """Стиль документа"""
        pass

    @document_style.setter
    def document_style(self, value: str) -> None:
        """Стиль документа"""
        pass

    @abstractmethod
    def get_document_style(self) -> Styles:
        """
        Получение стиля документа

        Returns:
            Styles: Стиль документа
        """
        pass

    @abstractmethod
    def apply_style(self, target: Self | Run | Paragraph, style: Any) -> None:
        """
        Применение стиля к передаваемому объекту

        Args:
            target (ABCDocument | Run | Paragraph): Цель применения стиля
            style (Any): Применяемый стиль
        """
        pass

    @abstractmethod
    def create_document(self, file_name: str, path: str | Path | None) -> None:
        """
        Создание документа

        Args:
            file_name (str): Наименование документа
            path (str | Path | None): Путь формирования документа
        """
        pass

    @abstractmethod
    def load_document(self) -> None:
        """Загрузка уже имеющегося документа"""
        pass

    @abstractmethod
    def save_document(self) -> None:
        """Сохранение документа"""
        pass

    @abstractmethod
    def run_instruction(self) -> None:
        """
        Запуск инструкции формирования документа
        """
        pass

    @property
    @abstractmethod
    def creation_instruction(self) -> Callable:
        """Функция для формирования документа"""
        pass

    @creation_instruction.setter
    def creation_instruction(self, value: Callable) -> None:
        """Функция для формирования документа"""
        pass

    @property
    @abstractmethod
    def instruction_kwargs(self) -> dict[str | Any] | None:
        """Аргументы инструкции"""
        pass

    @instruction_kwargs.setter
    def instruction_kwargs(self, value: dict[str | Any] | None) -> None:
        pass
