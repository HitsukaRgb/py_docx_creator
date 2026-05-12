import copy
import pprint
from dataclasses import dataclass, asdict
from typing import Any, Type

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

from py_docx_creator.AbstractClasses import DocumentStyles, FontNames, ParagraphStyle, TextStyle
from py_docx_creator.CoreClasses import CoreDocumentStyle, CorePageStyle, \
    CoreTextStyle, CoreDocumentWriter, AlignParagraph, CoreParagraphStyle, CoreStyleManager
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

class NormalDocumentStyle(CoreDocumentStyle):
    """Стандартный стиль документа"""

    def __init__(self):
        super().__init__()
        self.document_style = DocumentStyles.Normal.value


@dataclass
class MainPageStyle(CorePageStyle):
    """Основной формат страницы"""
    top_margin: float | None = 15.0
    bottom_margin: float | None = 10.0
    left_margin: float | None = 75.0
    right_margin: float | None = 75.0


@dataclass
class MainParagraphStyle(CoreParagraphStyle):
    """Стиль основного текста"""
    alignment: AlignParagraph | None = AlignParagraph.JUSTIFY.value
    space_after: float | None = 0.0
    left_indent: float | None = -0.5
    right_indent: float | None = -0.5
    line_spacing: float | None = 1.15
    first_line_indent: float | None = 20



@dataclass
class HeaderParagraphStyle(CoreParagraphStyle):
    """Стиль для заголовков """
    alignment: AlignParagraph = AlignParagraph.CENTER.value
    left_indent: float = -0.5
    right_indent: float = -0.5


class MainDocumentWriter(CoreDocumentWriter):
    def __init__(self):
        super().__init__()


@dataclass
class MainTextStyle(CoreTextStyle):
    """Основной стиль текста"""
    size: float = 10.0
    name: str = FontNames.TimesNewRoman.value
    bold: bool = False


@dataclass
class HeaderTextStyle(CoreTextStyle):
    size: float = 12.0
    name: str = FontNames.TimesNewRoman.value
    bold: bool = True


class FastWriter(CoreDocumentWriter):
    """Класс для быстрой записи в документ"""

    @classmethod
    def write(cls, document: Any, text: str, paragraph_style: Any, text_style: Any,
              size: float | None = None,
              bold: bool | None = None,
              italic: bool | None = None,
              underline: bool | None = None,
              space_after: float | None = None,
              alignment: Any | None = None,
              first_line_indent: float | None = None,
              with_leader: bool = False,
              leader_width: float = 6.8
              ) -> Any:

        # 1. Подготовка стилей (ваша логика без изменений)
        if any(val is not None for val in [bold, italic, underline, size, alignment, first_line_indent, space_after]):
            paragraph_style = copy.copy(paragraph_style())
            text_style = copy.copy(text_style())

            if bold is not None: text_style.bold = bold
            if italic is not None: text_style.italic = italic
            if underline is not None: text_style.underline = underline
            if size is not None: text_style.size = size
            if alignment is not None: paragraph_style.alignment = alignment.value
            if first_line_indent is not None: paragraph_style.first_line_indent = first_line_indent
            if space_after is not None: paragraph_style.space_after = space_after

        if hasattr(document, 'add_paragraph'):
            paragraph = cls.add_paragraph_to_document(document)
        else:
            paragraph = document

        if with_leader:
            text = f"{text}\t"
            tab_stops = paragraph.paragraph_format.tab_stops
            tab_stops.clear()
            tab_stops.add_tab_stop(
                Inches(leader_width),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
                leader=WD_TAB_LEADER.LINES
            )

        CoreStyleManager.PARAGRAPH_STYLE_MANAGER.apply_style(paragraph, paragraph_style)

        run = cls.add_run_to_paragraph(paragraph, text)
        CoreStyleManager.TEXT_STYLE_MANAGER.apply_style(run, text_style)

        return paragraph


